# -*- coding: utf-8 -*-
import sys
import os
import datetime
import re
from PIL import Image
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from PyQt5.QtWidgets import (
    QApplication, QWidget, QLabel, QLineEdit, QPushButton, QTextEdit,
    QVBoxLayout, QHBoxLayout, QFileDialog, QCheckBox, QSpinBox, QGroupBox,
    QMessageBox, QComboBox, QListWidget, QListWidgetItem, QScrollArea
)
from PyQt5.QtGui import QPixmap, QScreen, QGuiApplication
from PyQt5.QtCore import Qt, QRect
import logging
import ctypes # For admin check on Windows and Taskbar control

# --- Attempt to import keyboard and handle potential ImportError ---
try:
    import keyboard
    KEYBOARD_AVAILABLE = True
except ImportError:
    KEYBOARD_AVAILABLE = False
    print("WARNING: 'keyboard' library not found. Hotkey functionality will be disabled.")
    print("Install it using: pip install keyboard")

# --- Attempt to import optional libraries and handle potential ImportError ---
try:
    from docx2pdf import convert
    DOCX2PDF_AVAILABLE = True
except ImportError:
    DOCX2PDF_AVAILABLE = False
    print("WARNING: 'docx2pdf' library not found. PDF conversion will be disabled.")
    print("Install it using: pip install docx2pdf")
    # On Windows, docx2pdf often requires pywin32: pip install pywin32

try:
    import openpyxl
    from openpyxl.drawing.image import Image as ExcelImage
    OPENPYXL_AVAILABLE = True
except ImportError:
    OPENPYXL_AVAILABLE = False
    print("WARNING: 'openpyxl' library not found. Excel generation will be disabled.")
    print("Install it using: pip install openpyxl")


# --- Constants for Taskbar Control (Windows specific) ---
FindWindow = ctypes.windll.user32.FindWindowW
ShowWindow = ctypes.windll.user32.ShowWindow
SW_HIDE = 0
SW_SHOW = 5

# --- Helper function to check for Admin privileges on Windows ---
def is_admin():
    if os.name == 'nt':
        try:
            return ctypes.windll.shell32.IsUserAnAdmin()
        except:
            return False
    else:
        # Check if running as root on Unix-like systems
        try:
            return (os.geteuid() == 0)
        except AttributeError:
             return False # Assume not root if geteuid doesn't exist

# --- Helper function to control Taskbar visibility (Windows specific) ---
def set_taskbar_visibility(visible):
    # This function targets the main system taskbar ("Shell_TrayWnd")
    # It should work regardless of the number of monitors in standard Windows setups.
    hwnd = FindWindow("Shell_TrayWnd", None)
    if hwnd:
        state = SW_SHOW if visible else SW_HIDE
        ShowWindow(hwnd, state)
        # Also handle the Start button orb window if present
        hwnd_start = FindWindow("Button", "Start") # May need adjustment on different Windows versions
        if hwnd_start:
            ShowWindow(hwnd_start, state)
        logging.info(f"Taskbar visibility set to {'shown' if visible else 'hidden'}.")
        return True
    else:
        logging.warning("Could not find Taskbar window (Shell_TrayWnd) to change visibility.")
        return False


class ScreenshotApp(QWidget):
    def __init__(self):
        super().__init__()
        self.setWindowTitle("Screenshot Tool")
        self.screenshot_count = 1
        self.capture_enabled = False
        self.default_key = 'home'
        self.hotkey = self.default_key # Initialize hotkey
        self.registered_hotkey = None # Track successfully registered hotkey
        self.doc_path = None
        self.doc = None
        self.captured_data = []
        self.captured_images = []
        self.delete_images_after_save = False
        self.selected_monitors = []
        self.capture_mode = "single"
        self.excel_path = None
        self.taskbar_hidden = False # Track taskbar state

        # --- Check for Admin privileges early ---
        if not is_admin() and os.name == 'nt':
             print("WARNING: Script may not have administrator privileges. Hotkey registration might fail.")
             self.show_admin_warning = True
        else:
            self.show_admin_warning = False

        # Configure logging
        log_file = 'screenshot_app.log'
        try:
            # Try to remove old log file, ignore errors if it's locked etc.
            if os.path.exists(log_file):
                try:
                    os.remove(log_file)
                except OSError as remove_err:
                    print(f"Warning: Could not remove old log file '{log_file}': {remove_err}")

            logging.basicConfig(filename=log_file, level=logging.INFO,
                                format='%(asctime)s - %(levelname)s - %(message)s',
                                filemode='a') # Use 'a' (append) mode
            logging.info("-------------------- Application Started --------------------")
            logging.info(f"Admin/Root privileges: {is_admin()}")
            logging.info(f"Keyboard library available: {KEYBOARD_AVAILABLE}")
            logging.info(f"docx2pdf library available: {DOCX2PDF_AVAILABLE}")
            logging.info(f"openpyxl library available: {OPENPYXL_AVAILABLE}")
        except Exception as e:
            print(f"CRITICAL: Error setting up logging: {e}")

        self.init_ui()

        # --- Show Admin Warning Dialog AFTER UI is initialized ---
        if self.show_admin_warning:
              QMessageBox.warning(self, "Permissions Warning",
                                  "The script is not running as administrator.\n"
                                  "Registering global hotkeys might fail.\n"
                                  "Please restart with 'Run as administrator'.")

    def init_ui(self):
        # --- UI Setup Code ---
        main_layout = QVBoxLayout()
        scroll_area = QScrollArea()
        scroll_area.setWidgetResizable(True)
        scroll_content = QWidget()
        layout = QVBoxLayout(scroll_content)

        # Settings Group
        basic_group = QGroupBox("Settings")
        basic_group.setStyleSheet("color: black;font-weight: bold;")
        basic_layout = QVBoxLayout()
        basic_layout.addWidget(QLabel("Test Case Name:"))
        self.test_case_input = QLineEdit()
        self.test_case_input.setPlaceholderText("Enter Case Name (Default: Evidence)")
        basic_layout.addWidget(self.test_case_input)
        basic_layout.addWidget(QLabel("Document Version (e.g., v1):"))
        self.version_input = QLineEdit("v1")
        basic_layout.addWidget(self.version_input)
        basic_layout.addWidget(QLabel("Screenshot Key:"))
        self.hotkey_input = QLineEdit(self.default_key)
        if not KEYBOARD_AVAILABLE:
            self.hotkey_input.setEnabled(False)
            self.hotkey_input.setToolTip("Requires 'keyboard' library")
        basic_layout.addWidget(self.hotkey_input)
        basic_group.setLayout(basic_layout)
        layout.addWidget(basic_group)

        # Monitor Selection Group
        monitor_group = QGroupBox("Monitor Selection")
        monitor_group.setStyleSheet("color: black;font-weight: bold;")
        monitor_layout = QVBoxLayout()
        self.monitor_mode_combo = QComboBox()
        self.monitor_mode_combo.addItem("Single Monitor")
        self.monitor_mode_combo.addItem("Capture All Monitors (Stitched)")
        self.monitor_mode_combo.addItem("Select Multiple Monitors")
        self.monitor_mode_combo.currentIndexChanged.connect(self.monitor_mode_changed)
        monitor_layout.addWidget(self.monitor_mode_combo)
        self.single_monitor_combo = QComboBox()
        self.populate_single_monitor_combo()
        monitor_layout.addWidget(self.single_monitor_combo)
        self.multiple_monitor_list = QListWidget()
        self.populate_multiple_monitor_list()
        monitor_layout.addWidget(self.multiple_monitor_list)
        self.update_monitor_visibility()
        monitor_group.setLayout(monitor_layout)
        layout.addWidget(monitor_group)

        # Output Group
        output_group = QGroupBox("Output Settings")
        output_group.setStyleSheet("color: black;font-weight: bold;")
        output_layout = QVBoxLayout()
        output_layout.addWidget(QLabel("Output Folder:"))
        folder_layout = QHBoxLayout()
        self.folder_input = QLineEdit()
        self.folder_input.setPlaceholderText("Select folder to save documents")
        browse_button = QPushButton("Browse...")
        browse_button.clicked.connect(self.browse_folder)
        folder_layout.addWidget(self.folder_input)
        folder_layout.addWidget(browse_button)
        output_layout.addLayout(folder_layout)
        self.timestamp_checkbox = QCheckBox("Add Timestamp to Description")
        self.increment_checkbox = QCheckBox("Enable Auto-increment Count")
        self.increment_spin = QSpinBox()
        self.increment_spin.setValue(1); self.increment_spin.setMinimum(1)
        self.delete_checkbox = QCheckBox("Delete Images After Save")
        self.generate_excel_checkbox = QCheckBox("Generate Excel Document with Images")
        if not OPENPYXL_AVAILABLE:
             self.generate_excel_checkbox.setEnabled(False)
             self.generate_excel_checkbox.setToolTip("Requires 'openpyxl' library")

        # --- Add Taskbar Control Checkbox ---
        self.taskbar_checkbox = QCheckBox("Hide Taskbar During Capture")
        if os.name != 'nt': # Only enable on Windows
            self.taskbar_checkbox.setEnabled(False)
            self.taskbar_checkbox.setToolTip("Taskbar control only available on Windows")
        else:
             self.taskbar_checkbox.toggled.connect(self.toggle_taskbar_on_check)

        output_layout.addWidget(self.timestamp_checkbox)
        output_layout.addWidget(self.increment_checkbox)
        inc_layout = QHBoxLayout()
        inc_layout.addWidget(QLabel("Increment By:")); inc_layout.addWidget(self.increment_spin)
        output_layout.addLayout(inc_layout)
        output_layout.addWidget(self.delete_checkbox);
        output_layout.addWidget(self.generate_excel_checkbox)
        output_layout.addWidget(self.taskbar_checkbox) # Add the new checkbox here
        output_group.setLayout(output_layout)
        layout.addWidget(output_group)

        # Description & Preview
        layout.addWidget(QLabel("Screenshot Description:"))
        self.description_input = QLineEdit()
        layout.addWidget(self.description_input)
        layout.addWidget(QLabel("Screenshot Preview:"))
        self.preview_label = QLabel("Preview will appear here")
        self.preview_label.setAlignment(Qt.AlignCenter)
        self.preview_label.setFixedSize(300, 200)
        self.preview_label.setStyleSheet("border: 1px solid black; background-color: #f0f0f0;")
        layout.addWidget(self.preview_label)

        # Buttons
        button_layout = QHBoxLayout()
        self.start_button = QPushButton("Start New Capture");
        self.append_button = QPushButton("Append to Existing")
        self.stop_button = QPushButton("End Capture & Save");
        self.convert_pdf_button = QPushButton("Convert Doc to PDF")
        if not DOCX2PDF_AVAILABLE:
             self.convert_pdf_button.setEnabled(False);
             self.convert_pdf_button.setToolTip("Requires 'docx2pdf' library (pip install docx2pdf pywin32)") # Updated tooltip
             logging.info("Init UI: PDF button disabled (docx2pdf not available).") # Log reason
        else:
             self.convert_pdf_button.setEnabled(False); # Start disabled even if library exists
             logging.info("Init UI: PDF button initially disabled (awaiting successful save).")

        self.start_button.clicked.connect(self.start_new_capture); self.append_button.clicked.connect(self.append_to_existing)
        self.stop_button.clicked.connect(self.stop_capture);
        self.convert_pdf_button.clicked.connect(self.convert_to_pdf)
        self.stop_button.setEnabled(False)

        button_layout.addWidget(self.start_button);
        button_layout.addWidget(self.append_button)
        button_layout.addWidget(self.stop_button); button_layout.addWidget(self.convert_pdf_button)
        layout.addLayout(button_layout)

        # Status
        self.status_label = QLabel("Status: Ready")
        if not KEYBOARD_AVAILABLE: self.status_label.setText("Status: Ready (Hotkey disabled - install 'keyboard')")
        layout.addWidget(self.status_label)

        # Final layout setup
        scroll_content.setLayout(layout)
        scroll_area.setWidget(scroll_content)
        main_layout.addWidget(scroll_area)
        self.setLayout(main_layout)
        self.resize(650, 750) # Increased height slightly for new checkbox
        # --- End UI Setup Code ---

    def populate_single_monitor_combo(self):
        self.single_monitor_combo.clear()
        screens = QGuiApplication.screens()
        if not screens:
             logging.warning("No screens detected.")
             self.single_monitor_combo.addItem("No Monitors Found")
             return
        for i, screen in enumerate(screens):
            screen_name = screen.name() if screen.name() else f"Monitor {i + 1}"
            size = screen.size(); depth = screen.depth()
            primary = " (Primary)" if screen == QGuiApplication.primaryScreen() else ""
            self.single_monitor_combo.addItem(f"{screen_name} ({size.width()}x{size.height()}@{depth}bit){primary}", userData=i)
        logging.info(f"Populated single monitor combo box with {len(screens)} monitors.")

    def populate_multiple_monitor_list(self):
        self.multiple_monitor_list.clear()
        screens = QGuiApplication.screens()
        if not screens:
             logging.warning("No screens detected.")
             item = QListWidgetItem("No Monitors Found")
             item.setFlags(item.flags() & ~Qt.ItemIsUserCheckable)
             self.multiple_monitor_list.addItem(item)
             return
        for i, screen in enumerate(screens):
            screen_name = screen.name() if screen.name() else f"Monitor {i + 1}"
            size = screen.size(); depth = screen.depth()
            primary = " (Primary)" if screen == QGuiApplication.primaryScreen() else ""
            item = QListWidgetItem(f"{screen_name} ({size.width()}x{size.height()}@{depth}bit){primary}")
            item.setData(Qt.UserRole, i);
            item.setFlags(item.flags() | Qt.ItemIsUserCheckable); item.setCheckState(Qt.Unchecked)
            self.multiple_monitor_list.addItem(item)
        logging.info(f"Populated multiple monitor list with {len(screens)} monitors.")

    def update_monitor_visibility(self):
        mode = self.monitor_mode_combo.currentText()
        is_single = "Single Monitor" in mode
        is_multiple = "Select Multiple" in mode
        is_all = "Capture All" in mode

        self.single_monitor_combo.setVisible(is_single)
        self.multiple_monitor_list.setVisible(is_multiple)

        if is_single: self.capture_mode = "single"
        elif is_multiple: self.capture_mode = "multiple"
        elif is_all: self.capture_mode = "all"
        logging.info(f"Monitor visibility updated. Mode: {self.capture_mode}")

    def monitor_mode_changed(self, index):
        self.update_monitor_visibility()

    def browse_folder(self):
        folder = QFileDialog.getExistingDirectory(self, "Select Output Folder")
        if folder:
              self.folder_input.setText(folder)
              logging.info(f"Output folder selected: {folder}")

    def validate_settings(self, is_new_capture=True):
        folder = self.folder_input.text().strip()
        if not folder or not os.path.isdir(folder):
            QMessageBox.warning(self, "Missing Folder", "Please select a valid output folder.")
            logging.warning("Validation failed: Output folder not set or invalid.")
            return False

        if KEYBOARD_AVAILABLE and not self.hotkey_input.text().strip():
             QMessageBox.warning(self, "Missing Hotkey", "Please enter a screenshot capture key.")
             logging.warning("Validation failed: Hotkey not set.")
             return False
        elif not KEYBOARD_AVAILABLE:
              logging.warning("Skipping hotkey validation ('keyboard' missing).")


        if is_new_capture:
            case_name = self.test_case_input.text().strip() or "Evidence"
            version = self.version_input.text().strip() or "v1"
            doc_path_check = os.path.join(folder, f"{case_name}_{version}.docx")
            if os.path.exists(doc_path_check):
                 reply = QMessageBox.question(self, 'File Exists',
                                              f"The file '{os.path.basename(doc_path_check)}' already exists.\nDo you want to overwrite it?",
                                              QMessageBox.Yes | QMessageBox.No, QMessageBox.No)
                 if reply == QMessageBox.No:
                      logging.warning("Validation failed: User chose not to overwrite existing file.")
                      return False
                 else:
                      logging.info(f"User chose to overwrite existing file: {doc_path_check}")

        logging.info("Settings validation successful.")
        return True


    def register_hotkey(self):
        if not KEYBOARD_AVAILABLE:
            logging.warning("Cannot register hotkey: 'keyboard' library not available.")
            self.status_label.setText("Status: Ready (Hotkey disabled - install 'keyboard')")
            return True # Allow continuing without hotkey

        try:
            # Attempt to remove any previously registered hotkey by this instance
            if hasattr(self, 'registered_hotkey') and self.registered_hotkey:
                 try:
                     keyboard.remove_hotkey(self.registered_hotkey)
                     logging.info(f"Removed previous hotkey: {self.registered_hotkey}")
                     self.registered_hotkey = None
                 except KeyError:
                     logging.warning(f"Tried removing hotkey '{self.registered_hotkey}', but it was not registered by the 'keyboard' library.")
                     self.registered_hotkey = None
                 except Exception as e:
                     logging.error(f"Error removing previous hotkey '{self.registered_hotkey}': {e}")
                     self.registered_hotkey = None # Assume it's gone

            self.hotkey = self.hotkey_input.text().strip().lower()
            if not self.hotkey:
                logging.error("Hotkey cannot be empty.")
                QMessageBox.critical(self, "Error", "Hotkey cannot be empty.")
                return False

            # Register the new hotkey
            keyboard.add_hotkey(self.hotkey, self.capture_screenshot)
            self.registered_hotkey = self.hotkey # Store the successfully registered key
            logging.info(f"Registered new hotkey: {self.registered_hotkey}")
            self.status_label.setText(f"Status: Ready (Hotkey: {self.registered_hotkey})")
            return True
        except ImportError as e:
             # This might happen if the backend (like _winkeyboard) fails to load
             logging.error(f"Hotkey registration ImportError: {e}")
             QMessageBox.critical(self, "Hotkey Error", f"Failed to initialize keyboard backend.\nError: {e}")
             self.registered_hotkey = None
             return False
        except Exception as e:
            # General catch-all, often related to permissions
            error_msg = f"Could not register hotkey '{self.hotkey}'.\nError: {e}\n\n"
            priv = is_admin()
            if os.name == 'nt' and not priv:
                error_msg += "**Run as Administrator is likely required.**\n"
            elif os.name != 'nt' and not priv:
                error_msg += "**Run with root privileges (sudo) is likely required.**\n"
            error_msg += "\nTry a different key combination or ensure you have sufficient privileges."
            logging.error(f"Error registering hotkey '{self.hotkey}': {e}. Admin/Root likely required.")
            QMessageBox.critical(self, "Hotkey Error", error_msg)
            self.registered_hotkey = None # Ensure no hotkey is stored if registration failed
            return False


    def unregister_hotkey(self):
        if not KEYBOARD_AVAILABLE: return
        hotkey_to_remove = getattr(self, 'registered_hotkey', None)
        if hotkey_to_remove:
            try:
                keyboard.remove_hotkey(hotkey_to_remove)
                logging.info(f"Unhooked hotkey: {hotkey_to_remove}")
            except KeyError:
                # This can happen if the key wasn't actually registered or was removed externally
                logging.warning(f"Tried to unhook hotkey '{hotkey_to_remove}', but it was not found by the 'keyboard' library.")
            except Exception as e:
                logging.error(f"Error removing hotkey '{hotkey_to_remove}': {e}")
            finally:
                self.registered_hotkey = None # Clear our record regardless
        else:
            logging.info("Unhook hotkey called, but no hotkey was recorded as registered.")


    def start_new_capture(self):
        if not self.validate_settings(is_new_capture=True): return
        if not self.register_hotkey(): return # register_hotkey handles messages on failure

        try:
            self.capture_enabled = True; self.screenshot_count = 1
            folder = self.folder_input.text().strip(); case_name = self.test_case_input.text().strip() or "Evidence"; version = self.version_input.text().strip() or "v1"
            base_filename = f"{case_name}_{version}"; self.doc_path = os.path.join(folder, f"{base_filename}.docx"); self.excel_path = os.path.join(folder, f"{base_filename}.xlsx")
            self.doc = Document() # Create new document object

            # Add Header
            if len(self.doc.sections) > 0:
                 header = self.doc.sections[0].header;
                 # Clear any existing content in the header
                 for para in header.paragraphs:
                     para._p.getparent().remove(para._p)
                 hp = header.add_paragraph();
                 hr = hp.add_run(f"Test Case: {case_name}"); hr.font.name = 'Arial'; hr.font.size = Pt(12); hr.bold = True;
                 hp.alignment = WD_ALIGN_PARAGRAPH.CENTER
            else:
                logging.warning("Cannot add header: Document has no sections.")

            self.captured_data = []; self.captured_images = []; self.delete_images_after_save = self.delete_checkbox.isChecked()

            current_hotkey = getattr(self, 'registered_hotkey', None);
            status_suffix = f"(Hotkey: {current_hotkey})" if current_hotkey else "(Hotkey Disabled)"
            self.status_label.setText(f"Status: Capture Started {status_suffix}");
            # Update button states
            self.stop_button.setEnabled(True); self.start_button.setEnabled(False); self.append_button.setEnabled(False)
            if DOCX2PDF_AVAILABLE: self.convert_pdf_button.setEnabled(False) # Ensure PDF button is off
            logging.info("Button States: Stop=Enabled, Start=Disabled, Append=Disabled, PDF=Disabled")

            # --- Hide Taskbar if checked ---
            self.taskbar_hidden = False # Reset flag
            if os.name == 'nt' and self.taskbar_checkbox.isChecked():
                if set_taskbar_visibility(False): # Hide
                    self.taskbar_hidden = True # Track state only if successful
                    logging.info("Taskbar hidden for capture.")
                else:
                    logging.warning("Attempted to hide taskbar, but failed (set_taskbar_visibility returned False).")
                    # Optionally notify user? For now, just log it.


            logging.info(f"Started new capture. Doc: {self.doc_path}, Excel: {self.excel_path}, Del Img: {self.delete_images_after_save}, Taskbar Hidden: {self.taskbar_hidden}, Hotkey: {current_hotkey if current_hotkey else 'N/A'}")

        except Exception as e:
            logging.error(f"Error during start_new_capture: {e}", exc_info=True)
            QMessageBox.critical(self, "Error", f"An unexpected error occurred starting the new capture:\n{e}")
            self.capture_enabled = False
            # Reset UI elements potentially affected
            self.stop_button.setEnabled(False)
            self.start_button.setEnabled(True)
            self.append_button.setEnabled(True)
            if DOCX2PDF_AVAILABLE: self.convert_pdf_button.setEnabled(False)
            self.unregister_hotkey() # Clean up hotkey if start failed


    def append_to_existing(self):
         options = QFileDialog.Options()
         # Use user's Documents folder as a starting point if possible
         default_dir = os.path.expanduser("~/Documents")
         if not os.path.isdir(default_dir): default_dir = "" # Fallback if Documents doesn't exist

         file_path, _ = QFileDialog.getOpenFileName(self, "Select Existing Word Document to Append To", default_dir, "Word Documents (*.docx)", options=options)
         if not file_path:
             logging.info("Append operation cancelled by user.")
             return # User cancelled dialog

         if not self.register_hotkey(): return # register_hotkey handles messages on failure

         try:
             self.doc_path = file_path
             # Attempt to open the document
             try:
                 self.doc = Document(self.doc_path)
             except Exception as doc_open_e:
                 logging.error(f"Failed to open selected Word document '{self.doc_path}': {doc_open_e}", exc_info=True)
                 QMessageBox.critical(self, "File Error", f"Could not open the selected Word document.\nIt might be corrupted, password-protected, or not a valid .docx file.\n\nError: {doc_open_e}")
                 self.unregister_hotkey() # Unregister hotkey if file fails
                 return

             # Find the last screenshot number to continue incrementing
             max_screenshot_num = 0;
             pattern = re.compile(r"Screenshot\s+(\d+)", re.IGNORECASE) # Regex to find "Screenshot N"
             if self.doc.paragraphs: # Check if document has paragraphs
                 # Iterate backwards for potentially faster finding of the last entry
                 for p in reversed(self.doc.paragraphs):
                     match = pattern.search(p.text);
                     if match:
                         try:
                             num = int(match.group(1))
                             max_screenshot_num = max(max_screenshot_num, num)
                             # Optimization: If we find one, we can potentially break
                             # if numbers are strictly sequential, but let's check all
                             # just in case they are out of order.
                             # break # Uncomment if strict sequential numbering is assumed
                         except ValueError:
                             pass # Ignore if the number part isn't a valid integer
             self.screenshot_count = max_screenshot_num + 1;
             logging.info(f"Determined starting screenshot number for append: {self.screenshot_count}")

             self.capture_enabled = True;
             self.captured_data = []; # Reset lists for new captures in this session
             self.captured_images = [];
             self.delete_images_after_save = self.delete_checkbox.isChecked()

             current_hotkey = getattr(self, 'registered_hotkey', None);
             status_suffix = f"(Hotkey: {current_hotkey})" if current_hotkey else "(Hotkey Disabled)"
             self.status_label.setText(f"Status: Appending to {os.path.basename(self.doc_path)} {status_suffix}");
             # Update button states
             self.stop_button.setEnabled(True); self.start_button.setEnabled(False); self.append_button.setEnabled(False)
             if DOCX2PDF_AVAILABLE: self.convert_pdf_button.setEnabled(False) # Ensure PDF button is off
             logging.info("Button States: Stop=Enabled, Start=Disabled, Append=Disabled, PDF=Disabled")


             # Set output folder to the same directory as the selected file
             self.folder_input.setText(os.path.dirname(self.doc_path));
             # Define Excel path based on appended document name
             base_filename = os.path.splitext(os.path.basename(self.doc_path))[0];
             self.excel_path = os.path.join(os.path.dirname(self.doc_path), f"{base_filename}_appended_data.xlsx") # Changed Excel name slightly

             # --- Hide Taskbar if checked ---
             self.taskbar_hidden = False # Reset flag
             if os.name == 'nt' and self.taskbar_checkbox.isChecked():
                 if set_taskbar_visibility(False): # Hide
                     self.taskbar_hidden = True # Track state only if successful
                     logging.info("Taskbar hidden for capture.")
                 else:
                      logging.warning("Attempted to hide taskbar, but failed (set_taskbar_visibility returned False).")


             logging.info(f"Ready to append to: {self.doc_path}, Start Count: {self.screenshot_count}, Del Img: {self.delete_images_after_save}, Taskbar Hidden: {self.taskbar_hidden}, Excel Path: {self.excel_path}, Hotkey: {current_hotkey if current_hotkey else 'N/A'}")

         except Exception as e:
             logging.error(f"Error during append_to_existing setup: {e}", exc_info=True)
             QMessageBox.critical(self, "Append Error", f"An unexpected error occurred setting up the append operation:\n{e}")
             self.doc_path = None; self.doc = None; self.capture_enabled = False;
             # Reset UI
             self.stop_button.setEnabled(False)
             self.start_button.setEnabled(True)
             self.append_button.setEnabled(True)
             if DOCX2PDF_AVAILABLE: self.convert_pdf_button.setEnabled(False)
             self.unregister_hotkey() # Clean up hotkey


    def stop_capture(self):
        if not self.capture_enabled:
            logging.warning("Stop capture called, but capture was not enabled.")
            return # Nothing to do if not capturing

        try:
            logging.info("Stop capture initiated.")
            self.capture_enabled = False;
            self.unregister_hotkey() # Always try to unregister the hotkey

            # --- Restore Taskbar if it was hidden by this app ---
            if self.taskbar_hidden and os.name == 'nt':
                if set_taskbar_visibility(True): # Show the taskbar
                    logging.info("Taskbar restored successfully.")
                else:
                    # Log failure, but don't block rest of stop process
                    logging.warning("Failed to restore taskbar (set_taskbar_visibility returned False).")
                self.taskbar_hidden = False # Reset flag regardless of success/failure

            pdf_enable = False # Default to disabled
            save_successful = False # Track save status

            if self.doc and self.doc_path:
                logging.info(f"Attempting to save Word document to: {self.doc_path}")
                try:
                    # --- Attempt to save the document ---
                    self.doc.save(self.doc_path);
                    save_successful = True # Mark save as successful
                    status_msg = f"Capture complete. Word saved: {os.path.basename(self.doc_path)}";
                    logging.info(f"Word document saved successfully: {self.doc_path}")

                    # --- Attempt Excel generation if requested and possible ---
                    if self.generate_excel_checkbox.isChecked():
                        if OPENPYXL_AVAILABLE and self.captured_data:
                            logging.info("Generating Excel document...")
                            self.generate_excel() # This function handles its own errors/messages
                            # Check if excel_path exists after generation attempt
                            if self.excel_path and os.path.exists(self.excel_path):
                                status_msg += f" | Excel saved: {os.path.basename(self.excel_path)}"
                            else:
                                status_msg += " | Excel generation failed (see log)"
                                logging.warning("Excel generation was checked, but Excel file path is not set or file doesn't exist after generation attempt.")
                        elif not OPENPYXL_AVAILABLE:
                            status_msg += " | Excel skipped (library missing)"
                            logging.warning("Excel generation skipped - 'openpyxl' library not available.")
                        else: # Library available but no data captured
                             status_msg += " | Excel skipped (no data)"
                             logging.info("Excel generation skipped - no screenshot data was captured.")
                    else:
                        logging.info("Excel generation not requested.")

                    self.status_label.setText(status_msg); # Update status bar
                    pdf_enable = True # Enable PDF button only if save succeeded

                except PermissionError as e:
                    logging.error(f"PermissionError saving Word document '{self.doc_path}': {e}", exc_info=True)
                    QMessageBox.critical(self, "Save Error", f"Failed to save Word document due to permissions.\nCheck if the file is open elsewhere or if you have write access to the folder.\n\nError: {e}")
                    self.status_label.setText("Status: Error saving Word (Permission Denied)!");
                    pdf_enable = False # Ensure PDF button is disabled
                except Exception as e:
                    logging.error(f"Error saving Word document '{self.doc_path}': {e}", exc_info=True)
                    QMessageBox.critical(self, "Save Error", f"An unexpected error occurred saving the Word document:\n{e}")
                    self.status_label.setText("Status: Error saving Word document!");
                    pdf_enable = False # Ensure PDF button is disabled

            else:
                # This case shouldn't normally happen if capture was enabled, but handle it defensively.
                self.status_label.setText("Status: Capture ended. No active document reference.");
                logging.warning("Stop capture called, but self.doc or self.doc_path was None/Empty.")
                pdf_enable = False


            # --- Set PDF Button State ---
            # Condition: Library must be available AND Word save must have succeeded
            final_pdf_enable_state = DOCX2PDF_AVAILABLE and pdf_enable
            self.convert_pdf_button.setEnabled(final_pdf_enable_state)
            # Log the final decision for clarity
            if final_pdf_enable_state:
                logging.info("PDF conversion button ENABLED (Library available and Word save successful).")
            elif not DOCX2PDF_AVAILABLE:
                logging.info("PDF conversion button remains DISABLED (docx2pdf library not available).")
            else: # Library available, but save failed
                logging.info("PDF conversion button remains DISABLED (Word document save failed or was skipped).")


            # --- Cleanup Images if requested AND save was successful ---
            if self.delete_images_after_save and save_successful:
                logging.info("Cleanup requested and save was successful. Proceeding with image deletion.")
                self.cleanup_captured_images()
            elif self.delete_images_after_save and not save_successful:
                 logging.warning("Cleanup requested, but Word save failed. Skipping image deletion to preserve evidence.")
            else:
                 logging.info("Image cleanup not requested.")

            # --- Reset main button states ---
            self.start_button.setEnabled(True);
            self.append_button.setEnabled(True);
            self.stop_button.setEnabled(False)
            logging.info("Button States Reset: Start=Enabled, Append=Enabled, Stop=Disabled")

        except Exception as e:
            # Catch-all for unexpected errors during the stop process itself
            logging.error(f"Critical error during stop_capture routine: {e}", exc_info=True)
            QMessageBox.critical(self, "Stop Capture Error", f"An unexpected error occurred during the stop process:\n{e}")
            self.status_label.setText("Status: Error during stop capture!");
            # Ensure UI is reset to a safe state
            self.start_button.setEnabled(True);
            self.append_button.setEnabled(True);
            self.stop_button.setEnabled(False)
            if DOCX2PDF_AVAILABLE: self.convert_pdf_button.setEnabled(False) # Ensure PDF button is off on error
            self.taskbar_hidden = False # Assume taskbar should be shown


    def convert_to_pdf(self):
        # 1. Check if library is available (should be redundant if button is enabled, but good practice)
        if not DOCX2PDF_AVAILABLE:
            QMessageBox.critical(self, "Conversion Failed", "'docx2pdf' library is not installed or could not be loaded.\nPlease install it (`pip install docx2pdf pywin32`).")
            logging.error("Convert to PDF called, but DOCX2PDF_AVAILABLE is False.")
            return

        # 2. Check if a valid Word document path exists
        if not self.doc_path or not os.path.exists(self.doc_path):
            self.status_label.setText("Status: Cannot convert - Word doc path invalid or file missing.");
            QMessageBox.warning(self, "No Word Document", "Cannot convert to PDF because the source Word document path is missing or the file doesn't exist.\nPlease ensure capture was stopped and saved successfully.")
            logging.warning(f"Convert to PDF called, but doc_path is invalid or file missing: '{self.doc_path}'")
            return

        # 3. Define PDF path and update status
        pdf_path = os.path.splitext(self.doc_path)[0] + ".pdf";
        self.status_label.setText(f"Status: Converting '{os.path.basename(self.doc_path)}' to PDF...");
        logging.info(f"Attempting to convert '{self.doc_path}' to '{pdf_path}'")
        QApplication.processEvents() # Keep UI responsive

        # 4. Attempt conversion
        try:
            convert(self.doc_path, pdf_path);
            # Success: Update status and notify user
            self.status_label.setText(f"Status: Converted to PDF: {os.path.basename(pdf_path)}");
            QMessageBox.information(self, "Conversion Successful", f"Successfully converted Word document to PDF:\n{pdf_path}");
            logging.info(f"Successfully converted '{self.doc_path}' to '{pdf_path}'")
        except Exception as e:
            # Failure: Log detailed error and notify user
            err_str = str(e).lower()
            logging.error(f"Failed to convert '{self.doc_path}' to PDF: {e}", exc_info=True)

            # Provide more specific error messages based on common issues
            if any(sub in err_str for sub in ["pywintypes.com_error", "dispatch", "win32com", "word", "(-214", "rpc_server"]):
                msg = (f"PDF conversion failed. This usually requires Microsoft Word to be installed and accessible by the script.\n"
                       f"Things to check:\n"
                       f"- Is MS Word installed?\n"
                       f"- Is the Word application closed?\n"
                       f"- Are there any Word dialog boxes open?\n"
                       f"- Try running this script as Administrator.\n\n"
                       f"Error details: {e}")
            elif "permission denied" in err_str:
                msg = (f"PDF conversion failed due to a Permission Denied error.\n"
                       f"Things to check:\n"
                       f"- Do you have write permissions in the output folder ('{os.path.dirname(pdf_path)}')?\n"
                       f"- Is the source Word file ('{os.path.basename(self.doc_path)}') currently open?\n"
                       f"- Is the target PDF file ('{os.path.basename(pdf_path)}') currently open?\n\n"
                       f"Error details: {e}")
            elif "read-only" in err_str: # Check for read-only explicitly
                 msg = (f"PDF conversion failed, possibly because the source Word document is read-only or locked.\n"
                        f"Please check the file properties of:\n'{self.doc_path}'\n\n"
                        f"Error details: {e}")
            else:
                # Generic error message
                msg = f"An unexpected error occurred during PDF conversion:\n{e}"

            self.status_label.setText("Status: PDF conversion failed!");
            QMessageBox.critical(self, "Conversion Failed", msg)


    def cleanup_captured_images(self):
        if not self.captured_images:
            logging.info("Cleanup: No captured image paths recorded. Nothing to delete.")
            return

        deleted_count, error_count = 0, 0;
        logging.info(f"Cleanup: Starting deletion process for {len(self.captured_images)} recorded image paths...")
        # Iterate over a copy of the list, as we modify the original list inside the loop
        images_to_delete = list(self.captured_images)
        for img_path in images_to_delete:
            try:
                if img_path and os.path.exists(img_path): # Check path exists
                    os.remove(img_path);
                    deleted_count += 1;
                    logging.info(f"Deleted image file: {img_path}")
                    if img_path in self.captured_images: self.captured_images.remove(img_path) # Remove from list
                elif img_path:
                    logging.warning(f"Image file not found for deletion (already deleted or path incorrect?): {img_path}")
                    if img_path in self.captured_images: self.captured_images.remove(img_path) # Remove missing path from list
                else:
                     logging.warning("Encountered a None or empty image path in cleanup list.")
                     if img_path in self.captured_images: self.captured_images.remove(img_path) # Remove bad entry

            except PermissionError as pe:
                 logging.error(f"PermissionError deleting image file {img_path}: {pe}")
                 error_count += 1
                 # Do not remove from list on permission error, maybe retry later? For now, keep it.
            except Exception as e:
                logging.error(f"Error deleting image file {img_path}: {e}")
                error_count += 1
                # Remove from list even if error occurred to avoid repeated attempts? Maybe not desirable.
                # For now, we leave it in the list if deletion fails for reasons other than NotFound/None.

        result_msg = f"Cleanup complete. Successfully deleted: {deleted_count}, Errors: {error_count}. Remaining in list: {len(self.captured_images)}";
        logging.info(result_msg)
        if error_count > 0:
            QMessageBox.warning(self, "Cleanup Issues", f"Could not delete {error_count} image file(s). Please check the application log ('screenshot_app.log') for details.")


    def generate_excel(self):
        # Pre-checks moved to stop_capture, but keep internal checks just in case
        if not OPENPYXL_AVAILABLE:
            QMessageBox.critical(self, "Excel Generation Failed", "'openpyxl' library is not installed.\nPlease install it (`pip install openpyxl`).")
            logging.error("Generate Excel called, but OPENPYXL_AVAILABLE is False.")
            return
        if not self.excel_path:
            logging.error("Generate Excel called, but self.excel_path is not set.")
            QMessageBox.warning(self, "Excel Error", "Cannot generate Excel because the output file path has not been set.")
            return
        if not self.captured_data:
            logging.warning("Generate Excel called, but self.captured_data is empty.")
            QMessageBox.information(self, "No Data", "No screenshot data was captured to put into an Excel file.")
            return

        logging.info(f"Generating Excel file at: {self.excel_path}")
        try:
            workbook = openpyxl.Workbook();
            sheet = workbook.active;
            sheet.title = "Screenshots";
            # Add header row
            sheet.append(["Screenshot No.", "Description", "Image"])
            # Set column widths (adjust as needed)
            sheet.column_dimensions['A'].width = 15; # Screenshot number
            sheet.column_dimensions['B'].width = 40; # Description
            sheet.column_dimensions['C'].width = 50  # Image column (width doesn't directly control image size)

            # Freeze header row
            sheet.freeze_panes = 'A2'

            # Style for headers (optional)
            header_font = openpyxl.styles.Font(bold=True)
            header_alignment = openpyxl.styles.Alignment(horizontal='center')
            for cell in sheet["1:1"]:
                cell.font = header_font
                cell.alignment = header_alignment


            img_height_pixels = 200 # Approx desired image height in Excel


            for row_num, data in enumerate(self.captured_data, start=2): # Start from row 2
                # Add screenshot number and description
                cell_co = sheet.cell(row=row_num, column=1, value=data.get("co", "N/A"))
                cell_co.alignment = openpyxl.styles.Alignment(horizontal='center', vertical='center')

                cell_desc = sheet.cell(row=row_num, column=2, value=data.get("description", ""))
                cell_desc.alignment = openpyxl.styles.Alignment(vertical='top', wrap_text=True)

                img_path = data.get("image_path")
                if img_path and os.path.exists(img_path):
                    try:
                        # Insert image
                        img = ExcelImage(img_path)

                        # --- Resize image proportionally to fit target height ---
                        try:
                            with Image.open(img_path) as pil_img:
                                w_px, h_px = pil_img.size
                                if h_px > 0 and w_px > 0: # Avoid division by zero
                                    aspect = w_px / h_px
                                    img.height = img_height_pixels
                                    img.width = img_height_pixels * aspect
                                else: # Handle zero dimension images? Set default size.
                                    img.width=100
                                    img.height=100
                                    logging.warning(f"Image has zero dimension: {img_path}")
                        except Exception as pil_e:
                            logging.error(f"PIL error processing image {img_path} for resizing: {pil_e}")
                            # Use default size if PIL fails
                            img.width=150
                            img.height=100

                        # Anchor image to the cell in column C
                        cell_ref = sheet.cell(row=row_num, column=3).coordinate
                        img.anchor = cell_ref;
                        sheet.add_image(img);

                        # Adjust row height to fit the image (convert pixels to points)
                        # Excel uses points (1/72 inch), screen uses pixels (often 96 DPI)
                        row_h_pts = img.height * 72 / 96
                        sheet.row_dimensions[row_num].height = max(row_h_pts, 20) # Ensure a minimum height

                        logging.info(f"Added image '{os.path.basename(img_path)}' to Excel row {row_num}")

                    except Exception as img_e:
                        logging.error(f"Error adding image '{img_path}' to Excel row {row_num}: {img_e}", exc_info=True)
                        sheet.cell(row=row_num, column=3, value=f"[Error adding image: {img_e}]")
                elif img_path:
                    logging.warning(f"Image file not found, cannot add to Excel: {img_path}")
                    sheet.cell(row=row_num, column=3, value="[Image file not found]")
                else:
                    logging.warning(f"No image path provided for data in row {row_num}.")
                    sheet.cell(row=row_num, column=3, value="[No image path]")


            # Save the workbook
            workbook.save(self.excel_path);
            logging.info(f"Excel file generated successfully: {self.excel_path}")
            # Optional: Notify user success (already done in stop_capture status)
            # QMessageBox.information(self, "Excel Generated", f"Excel document generated successfully:\n{self.excel_path}");

        except PermissionError as e:
            error_message = f"Failed to save Excel file due to permissions.\nCheck if the file is open elsewhere or if you have write access.\nPath: {self.excel_path}\n\nError: {e}";
            # self.status_label.setText("Status: Error generating Excel (Permission Denied)."); # Status already updated in stop_capture
            QMessageBox.critical(self, "Excel Save Failed", error_message);
            logging.error(error_message, exc_info=True)
        except Exception as e:
            error_message = f"An unexpected error occurred generating the Excel file:\n{e}";
            # self.status_label.setText("Status: Error generating Excel."); # Status already updated in stop_capture
            QMessageBox.critical(self, "Excel Generation Failed", error_message);
            logging.error(f"Error generating Excel file '{self.excel_path}': {e}", exc_info=True)


    def capture_screenshot(self):
        # --- Pre-capture checks ---
        if not self.capture_enabled:
            logging.warning("Screenshot trigger ignored: Capture not currently enabled.")
            return
        if self.doc is None:
            logging.warning("Screenshot trigger ignored: No active Word document object.")
            # Should we stop capture here? Maybe just warn.
            QMessageBox.warning(self, "Capture Error", "Cannot capture screenshot: No active Word document. Please start or append first.")
            return

        folder = self.folder_input.text().strip()
        if not os.path.isdir(folder):
            logging.error(f"Output folder '{folder}' is invalid or does not exist. Stopping capture.")
            QMessageBox.critical(self, "Folder Error", f"The selected output folder is invalid or no longer exists:\n{folder}\n\nPlease select a valid folder and restart capture.")
            self.stop_capture(); # Stop the capture process fully
            return

        # --- Prepare filenames and descriptions ---
        case_name = self.test_case_input.text().strip() or "Evidence";
        description = self.description_input.text().strip(); # Get description at time of capture
        current_count_base = self.screenshot_count # Base number for this trigger event

        timestamp_str = ("_" + datetime.datetime.now().strftime("%Y%m%d_%H%M%S")) if self.timestamp_checkbox.isChecked() else ""
        # Sanitize description for use in filename (remove invalid chars)
        safe_description_part = re.sub(r'[\\/*?:"<>|]', "", description)[:50] # Limit length
        # Base filename structure
        filename_base_template = f"{case_name}_SS{current_count_base}_{safe_description_part}{timestamp_str}".replace(" ", "_")

        screens = QGuiApplication.screens()
        if not screens:
            logging.error("No screens detected by QGuiApplication.")
            QMessageBox.warning(self, "Screen Error", "Could not detect any screens. Cannot capture.")
            return

        # --- Perform Capture based on Mode ---
        try:
            if self.capture_mode == "single":
                selected_index = self.single_monitor_combo.currentData() # UserData holds the index
                if selected_index is not None and 0 <= selected_index < len(screens):
                    screen = screens[selected_index]; geometry = screen.geometry()
                    logging.info(f"Capturing single monitor: Index={selected_index}, Name='{screen.name()}', Geometry={geometry}")

                    # Use the specific screen's grabWindow for potentially better results on that monitor
                    image = screen.grabWindow(0, int(geometry.x()), int(geometry.y()), int(geometry.width()), int(geometry.height()))

                    if image.isNull():
                        raise ValueError(f"grabWindow returned a null pixmap for single monitor index {selected_index}.")

                    image_path = os.path.join(folder, f"{filename_base_template}_Monitor{selected_index + 1}.png")
                    word_description = f"Screenshot {current_count_base} (Monitor {selected_index + 1}): {description}"

                    # --- Save, Add to Word/Data, Update Preview ---
                    self.process_and_save_capture(image, image_path, word_description, current_count_base, description)

                else:
                    logging.error(f"Invalid monitor index selected for single capture: {selected_index}")
                    QMessageBox.warning(self, "Monitor Error", f"The selected monitor (Index: {selected_index}) is invalid. Please select a valid monitor.")
                    return # Don't proceed with capture

            elif self.capture_mode == "all":
                 # Calculate the bounding box of the entire virtual desktop
                 all_x = min(s.geometry().x() for s in screens);
                 all_y = min(s.geometry().y() for s in screens)
                 all_right = max(s.geometry().right() for s in screens);
                 all_bottom = max(s.geometry().bottom() for s in screens)
                 # Ensure width/height are at least 1
                 all_width = max(1, all_right - all_x + 1);
                 all_height = max(1, all_bottom - all_y + 1)
                 logging.info(f"Capturing all monitors stitched: Virtual Desktop Bounds=({all_x},{all_y}, W={all_width}, H={all_height})")

                 try:
                      # Grab the entire virtual desktop area using the primary screen's context
                      image = QGuiApplication.primaryScreen().grabWindow(0, int(all_x), int(all_y), int(all_width), int(all_height))
                      if image.isNull():
                          raise ValueError("grabWindow returned a null pixmap when capturing the 'all monitors' virtual desktop.")
                 except Exception as grab_e:
                      logging.error(f"Error during grabWindow for 'all monitors': {grab_e}", exc_info=True)
                      QMessageBox.critical(self, "Capture Error", f"Failed to capture the combined screen area for 'all monitors'.\nError: {grab_e}")
                      return # Stop if grab failed

                 image_path = os.path.join(folder, f"{filename_base_template}_AllMonitors.png");
                 word_description = f"Screenshot {current_count_base} (All Monitors): {description}";

                 # --- Save, Add to Word/Data, Update Preview ---
                 self.process_and_save_capture(image, image_path, word_description, current_count_base, description)


            elif self.capture_mode == "multiple":
                 selected_indices = []
                 for i in range(self.multiple_monitor_list.count()):
                     item = self.multiple_monitor_list.item(i)
                     # Check if item is checkable, checked, and has valid data
                     if item.flags() & Qt.ItemIsUserCheckable and item.checkState() == Qt.Checked:
                         monitor_index = item.data(Qt.UserRole)
                         if monitor_index is not None and 0 <= monitor_index < len(screens):
                             selected_indices.append(monitor_index)
                         else:
                             logging.warning(f"Ignoring checked item '{item.text()}' with invalid monitor index data: {monitor_index}")

                 if not selected_indices:
                     QMessageBox.warning(self,"Selection Error","Please check at least one valid monitor in the list to capture.")
                     return # Stop if no valid monitors selected

                 logging.info(f"Starting multi-monitor capture (Crop Method) for indices: {selected_indices}")

                 # --- Grab the entire virtual desktop ONCE ---
                 all_x = min(s.geometry().x() for s in screens); all_y = min(s.geometry().y() for s in screens)
                 all_right = max(s.geometry().right() for s in screens); all_bottom = max(s.geometry().bottom() for s in screens)
                 all_width = max(1, all_right - all_x + 1); all_height = max(1, all_bottom - all_y + 1)
                 logging.info(f"Multi-capture: Virtual desktop bounds=({all_x},{all_y}, W={all_width}, H={all_height})")

                 try:
                      virtual_desktop_pixmap = QGuiApplication.primaryScreen().grabWindow(0, int(all_x), int(all_y), int(all_width), int(all_height))
                      if virtual_desktop_pixmap.isNull():
                          raise ValueError("grabWindow returned a null pixmap for virtual desktop (multi-monitor base).")
                      logging.info(f"Grabbed virtual desktop pixmap successfully: Size=({virtual_desktop_pixmap.width()}x{virtual_desktop_pixmap.height()})")
                 except Exception as grab_e:
                      logging.error(f"FAILED to grab virtual desktop base image for multi-monitor crop: {grab_e}", exc_info=True)
                      QMessageBox.critical(self, "Capture Error", f"Failed the initial screen grab required for multi-monitor capture.\nError: {grab_e}")
                      return # Cannot proceed if the base grab failed

                 # --- Process each selected monitor ---
                 captured_paths_multi = []; success_count = 0; error_during_multi = False
                 last_successful_path = None

                 for i, monitor_index in enumerate(selected_indices):
                     logging.info(f"--- Processing selected monitor index: {monitor_index} ---")
                     # Re-validate index just in case screen list changed (unlikely but safe)
                     if 0 <= monitor_index < len(screens):
                         screen = screens[monitor_index]; geometry = screen.geometry()
                         logging.info(f"Target Monitor: Name='{screen.name()}', Absolute Geometry={geometry}")

                         # Calculate crop rectangle relative to the full virtual desktop grab
                         # Ensure coordinates are integers
                         crop_x = int(geometry.x() - all_x)
                         crop_y = int(geometry.y() - all_y)
                         crop_w = int(geometry.width())
                         crop_h = int(geometry.height())
                         crop_rect = QRect(crop_x, crop_y, crop_w, crop_h)
                         logging.info(f"Calculated Crop Rect relative to virtual grab: {crop_rect}")

                         # Generate filename specific to this monitor
                         multi_filename = f"{filename_base_template}_Monitor{monitor_index + 1}.png"
                         multi_image_path = os.path.join(folder, multi_filename)
                         logging.info(f"Target save path for this monitor: {multi_image_path}")

                         try:
                              # Crop the image from the full grab
                              image = virtual_desktop_pixmap.copy(crop_rect)
                              if image.isNull():
                                  raise ValueError(f"Cropped pixmap is null for monitor {monitor_index+1}. Crop Rect: {crop_rect}. Grabbed size: {virtual_desktop_pixmap.width()}x{virtual_desktop_pixmap.height()}")
                              logging.info(f"Cropped pixmap size: {image.width()}x{image.height()}")

                              # Define description for this specific monitor
                              multi_word_desc = f"Screenshot {current_count_base} (Monitor {monitor_index + 1} of Multiple): {description}"
                              sub_count_label = f"{current_count_base}-{success_count+1}" # Unique ID for Excel (e.g., 5-1, 5-2)

                              # --- Save, Add to Word/Data ---
                              # Note: Preview is updated only with the last successful image after the loop
                              save_success = image.save(multi_image_path, "PNG")
                              if not save_success:
                                  raise IOError(f"QPixmap.save() returned False when trying to save {multi_image_path}")
                              logging.info(f"Successfully saved cropped image: {multi_image_path}")

                              # Add to Word (potentially add page break before subsequent images in multi-capture)
                              self.add_to_word(multi_image_path, multi_word_desc, new_page=(success_count > 0))
                              # Add to internal lists
                              self.captured_images.append(multi_image_path)
                              self.captured_data.append({"co": sub_count_label, "description": f"{description} (Monitor {monitor_index+1})", "image_path": multi_image_path})

                              success_count += 1
                              last_successful_path = multi_image_path # Store path for preview

                         except Exception as multi_proc_e:
                              logging.error(f"Error cropping or saving for monitor index {monitor_index} (Path: {multi_image_path}): {multi_proc_e}", exc_info=True)
                              QMessageBox.warning(self, "Multi-Capture Error", f"Failed to process or save the screenshot for monitor {monitor_index+1}.\nCheck logs for details.\n\nError: {multi_proc_e}")
                              error_during_multi = True
                     else:
                         logging.warning(f"Invalid monitor index {monitor_index} encountered during multi-capture processing loop. Skipping.")
                         error_during_multi = True # Treat as an error condition

                 # --- Update Preview after processing all selected monitors ---
                 if last_successful_path:
                     self.update_preview(last_successful_path)
                     logging.info(f"Multi-monitor capture finished. Preview updated with: {os.path.basename(last_successful_path)}")
                 elif error_during_multi:
                      self.preview_label.setText("Preview Error\n(Multi-capture issues)");
                      self.preview_label.setStyleSheet("border: 1px solid red; color: red;")
                      logging.warning("Multi-monitor capture finished, but encountered errors. Preview shows error state.")
                 else: # No successes and no errors? (Shouldn't happen if selected_indices was populated)
                      self.preview_label.setText("No Images Captured")
                      self.preview_label.setStyleSheet("border: 1px solid black;")
                      logging.warning("Multi-monitor capture finished with no successful captures and no logged errors.")

            # --- Update status and increment base count ---
            # Increment base count regardless of errors in multi-mode, as the "event" happened.
            self.status_label.setText(f"Status: Screenshot Event {current_count_base} processed.")
            increment_value = self.increment_spin.value() if self.increment_checkbox.isChecked() else 1
            self.screenshot_count += increment_value
            logging.info(f"Incrementing screenshot count. Next base count will be: {self.screenshot_count}")


        except IOError as e: # Catch file saving errors here specifically if they escape process_and_save
            logging.error(f"IOError during screenshot save: {e}", exc_info=True)
            QMessageBox.critical(self, "Save Error", f"Failed to save the screenshot file.\nCheck folder permissions and disk space.\n\nError: {e}")
            # Don't increment count if save fails? Or do? Let's increment anyway.
        except Exception as e: # Catch all other unexpected errors during capture
            logging.error(f"Critical error during capture_screenshot routine: {e}", exc_info=True)
            QMessageBox.critical(self, "Capture Error", f"An unexpected error occurred during the capture process:\n{e}")
            # Don't increment count on critical failure? Or do? Let's increment.


    def process_and_save_capture(self, image_pixmap, image_path, word_description, count_for_excel, desc_for_excel):
        """Helper function to save image, add to Word, add data, and update preview."""
        try:
            logging.info(f"Attempting to save image to {image_path}...")
            save_success = image_pixmap.save(image_path, "PNG")
            if not save_success:
                raise IOError(f"QPixmap.save() returned False for path: {image_path}")
            logging.info(f"Successfully saved image: {image_path}")

            # Add to Word document
            self.add_to_word(image_path, word_description)

            # Add to internal data lists
            self.captured_images.append(image_path)
            self.captured_data.append({"co": count_for_excel, "description": desc_for_excel, "image_path": image_path})

            # Update the UI preview
            self.update_preview(image_path)

        except FileNotFoundError as fnf_e: # Specifically catch if image path is invalid before Word add
             logging.error(f"Image file not found when trying to add to Word: {image_path}. Error: {fnf_e}", exc_info=True)
             QMessageBox.warning(self, "Processing Error", f"Saved image file not found, cannot add to Word or update preview.\nPath: {image_path}")
             # Still add data entry? Maybe mark as error?
             self.captured_data.append({"co": count_for_excel, "description": f"[SAVE ERROR] {desc_for_excel}", "image_path": image_path})

        except IOError as io_e: # Catch save errors
            logging.error(f"IOError saving image file {image_path}: {io_e}", exc_info=True)
            QMessageBox.critical(self, "Save Error", f"Failed to save screenshot file.\nCheck folder permissions and disk space.\nPath: {image_path}\n\nError: {io_e}")
            # Add data entry marked with error
            self.captured_data.append({"co": count_for_excel, "description": f"[SAVE ERROR] {desc_for_excel}", "image_path": image_path})

        except Exception as e: # Catch other errors (e.g., Word processing)
            logging.error(f"Error processing capture (Image: {os.path.basename(image_path)}): {e}", exc_info=True)
            QMessageBox.warning(self, "Processing Error", f"An error occurred adding the screenshot to the document or updating data.\nImage: {os.path.basename(image_path)}\n\nError: {e}")
            # Add data entry marked with error
            self.captured_data.append({"co": count_for_excel, "description": f"[PROCESSING ERROR] {desc_for_excel}", "image_path": image_path})


    def add_to_word(self, image_path, description, new_page=False):
        if not self.doc:
            logging.error("Add to Word failed: Document object (self.doc) is None.")
            return
        try:
            # Add page break before the description if requested and not the first element
            if new_page and len(self.doc.paragraphs) > 1: # Check if paragraphs exist beyond initial section setup
                # Heuristic: Check if the last paragraph is not empty before adding break
                # This avoids breaks after section headers if nothing else was added yet.
                # A more robust check might be needed depending on document structure.
                 last_para = self.doc.paragraphs[-1]
                 if last_para.text.strip() or len(last_para.runs) > 0: # Check if last para has content
                      self.doc.add_page_break();
                      logging.info("Added page break before new content.")


            # Add the description paragraph
            p_desc = self.doc.add_paragraph();
            run = p_desc.add_run(description);
            # Style the description
            run.font.name = 'Arial'; run.font.size = Pt(10); run.italic = True

            # Add the image, scaled to fit page width
            section = self.doc.sections[-1]; # Get current section
            # Get page dimensions and margins (provide defaults if None)
            page_w = section.page_width if section.page_width else Inches(8.5)
            margin_l = section.left_margin if section.left_margin else Inches(1.0)
            margin_r = section.right_margin if section.right_margin else Inches(1.0)

            # Calculate available width for the image in EMUs (English Metric Units)
            # Add a small buffer (e.g., 98%) to prevent slight overflows
            available_width_emu = (page_w - margin_l - margin_r) * 0.98

            # Convert available width to inches (1 inch = 914400 EMUs)
            available_width_inches = available_width_emu / 914400.0

            if available_width_inches <= 0:
                # Fallback if calculation fails
                logging.warning(f"Calculated available page width is zero or negative ({available_width_inches:.2f} inches). Adding image '{os.path.basename(image_path)}' with default size.")
                pic_para = self.doc.add_paragraph() # Add image in its own paragraph for centering
                pic_run = pic_para.add_run()
                pic_run.add_picture(image_path)
                pic_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

            else:
                logging.info(f"Adding picture '{os.path.basename(image_path)}' scaled to fit available width ({available_width_inches:.2f} inches).")
                pic_para = self.doc.add_paragraph() # Add image in its own paragraph for centering
                pic_run = pic_para.add_run()
                # Add picture scaled to calculated width
                pic_run.add_picture(image_path, width=Inches(available_width_inches))
                pic_para.alignment = WD_ALIGN_PARAGRAPH.CENTER # Center the paragraph containing the image

            # Add a blank paragraph for spacing after the image
            self.doc.add_paragraph("")
            logging.info(f"Successfully added description and image '{os.path.basename(image_path)}' to Word.")

        except FileNotFoundError:
            # Handle case where image file doesn't exist when trying to add it
            logging.error(f"Add to Word failed: Image file not found at path: {image_path}")
            # Add an error message to the Word document itself
            self.doc.add_paragraph(f"[Error: Image file not found: {os.path.basename(image_path)}]")
        except Exception as e:
            # Handle other potential errors (e.g., issues with python-docx, invalid image format for Word)
            logging.error(f"Error adding content to Word document (Image: {os.path.basename(image_path)}): {e}", exc_info=True)
            # Add an error message to the Word document itself
            self.doc.add_paragraph(f"[Error adding content for image {os.path.basename(image_path)}: {e}]")


    def update_preview(self, image_path):
        if not image_path or not os.path.exists(image_path):
            self.preview_label.setText("Preview Error\n(File not found)");
            self.preview_label.setStyleSheet("border: 1px solid red; color: red;"); # Error style
            logging.warning(f"Preview update skipped: Image path invalid or file not found: '{image_path}'");
            return

        try:
            # Load the pixmap from the file
            pixmap = QPixmap(image_path);
            if pixmap.isNull():
                # This can happen if the file is corrupted or not a supported image format
                raise ValueError(f"Loaded pixmap is null. File might be corrupted or invalid format: {image_path}")

            # Scale the pixmap to fit the preview label while keeping aspect ratio
            pixmap_scaled = pixmap.scaled(self.preview_label.size(), Qt.KeepAspectRatio, Qt.SmoothTransformation)

            # Set the scaled pixmap on the label
            self.preview_label.setPixmap(pixmap_scaled);
            self.preview_label.setStyleSheet("border: 1px solid black;"); # Reset to normal style
            logging.info(f"Preview updated successfully with image: {os.path.basename(image_path)}")

        except Exception as e:
            # Handle errors during pixmap loading or scaling
            logging.error(f"Error updating preview with image {image_path}: {e}", exc_info=True);
            self.preview_label.setText("Preview Error\n(Load/Scale failed)");
            self.preview_label.setStyleSheet("border: 1px solid red; color: red;"); # Error style


    def toggle_taskbar_on_check(self, checked):
         # This allows manual toggling via checkbox, distinct from capture start/stop
         if os.name == 'nt':
             logging.info(f"Taskbar checkbox toggled to: {checked}")
             if checked:
                 # --- Try to Hide Taskbar ---
                 if not self.taskbar_hidden: # Only hide if not already hidden
                     if set_taskbar_visibility(False):
                         logging.info("Taskbar hidden manually via checkbox.")
                         # We set self.taskbar_hidden = True here IF we want the stop/close
                         # actions to automatically restore it even if hidden manually.
                         # Let's assume manual toggle = manual restore for now, UNLESS capture is active.
                         # If capture starts later, it will take over the state.
                         # self.taskbar_hidden = True # Decide if manual hide should be auto-restored
                     else:
                         logging.warning("Failed to hide taskbar via checkbox.")
                         QMessageBox.warning(self, "Taskbar Control Failed", "Could not hide the taskbar.")
                         self.taskbar_checkbox.setChecked(False) # Revert checkbox state on failure
                 else:
                      logging.info("Taskbar already hidden, checkbox toggle ignored.")
             else:
                 # --- Try to Show Taskbar ---
                 # Only explicitly show if we are SURE it's hidden.
                 # The self.taskbar_hidden flag tracks if WE hid it during capture.
                 # It's ambiguous if it was hidden manually or by another app.
                 # Safest might be to only attempt show if self.taskbar_hidden is True,
                 # OR maybe always attempt show when unchecked? Let's try always attempting.
                 if set_taskbar_visibility(True):
                     logging.info("Taskbar shown manually via checkbox.")
                     self.taskbar_hidden = False # Assume it's shown now
                 else:
                     logging.warning("Failed to show taskbar via checkbox.")
                     QMessageBox.warning(self, "Taskbar Control Failed", "Could not show the taskbar.")
                     # Should we re-check the box? Might be confusing. Leave as unchecked.


    def closeEvent(self, event):
        logging.info("-------------------- Application Closing --------------------")
        # --- Ensure Taskbar is restored on close IF we hid it ---
        if self.taskbar_hidden and os.name == 'nt':
            logging.info("Restoring taskbar visibility on application close...")
            set_taskbar_visibility(True) # Attempt to restore
            self.taskbar_hidden = False # Reset flag

        self.unregister_hotkey() # Attempt to unregister hotkey
        logging.shutdown() # Flush and close log file handlers
        event.accept() # Allow the window to close


if __name__ == "__main__":
    # Set AppUserModelID for Windows Taskbar icon behavior (optional but good practice)
    if os.name == 'nt':
         try:
             myappid = 'CompanyName.ProductName.ScreenshotTool.1.4'; # Unique ID
             ctypes.windll.shell32.SetCurrentProcessExplicitAppUserModelID(myappid)
         except Exception as e:
             print(f"Warning: Could not set AppUserModelID: {e}")

    app = QApplication(sys.argv)

    # --- Critical Dependency Check: PyQt5 ---
    try:
        # Test if basic Qt modules load, more specific than just the import line
        from PyQt5.QtWidgets import QApplication, QMessageBox
    except ImportError:
         # Use a simple Tkinter fallback for the error message if PyQt isn't even importable
         try:
             import tkinter as tk
             from tkinter import messagebox
             root = tk.Tk()
             root.withdraw() # Hide the main Tk window
             messagebox.showerror("Fatal Error",
                                  "Required library PyQt5 is missing or failed to load.\n"
                                  "Please install it using:\n"
                                  "pip install PyQt5")
             root.destroy()
         except ImportError:
             # Absolute fallback if neither Qt nor Tkinter is available
             print("FATAL ERROR: Required library PyQt5 is missing or failed to load.")
             print("Install using: pip install PyQt5")
         sys.exit(1) # Exit script

    # --- Informational Reminders about Optional Libraries ---
    if not KEYBOARD_AVAILABLE: print("Reminder: Global Hotkey feature disabled ('keyboard' library missing or failed to load).")
    if not DOCX2PDF_AVAILABLE: print("Reminder: PDF Conversion feature disabled ('docx2pdf' library missing or failed to load).")
    if not OPENPYXL_AVAILABLE: print("Reminder: Excel Generation feature disabled ('openpyxl' library missing or failed to load).")

    # --- Create and Show Main Window ---
    try:
        window = ScreenshotApp()
        window.show()
        sys.exit(app.exec_())
    except Exception as main_e:
         logging.error(f"Unhandled exception in main execution block: {main_e}", exc_info=True)
         QMessageBox.critical(None, "Application Crash", f"An unexpected error occurred and the application needs to close:\n{main_e}")
         # Ensure taskbar is restored if possible before crashing
         if getattr(window, 'taskbar_hidden', False) and os.name == 'nt':
              set_taskbar_visibility(True)
         sys.exit(1)
